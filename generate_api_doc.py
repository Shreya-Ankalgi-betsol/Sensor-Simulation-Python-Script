from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_to_doc(doc, title, columns, rows):
    """Helper to add a styled table"""
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = col
        # Style header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

# Create document
doc = Document()

# Title
title = doc.add_heading('BACKEND API DOCUMENTATION', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph('Version: 1.0.0 | Generated: March 30, 2026')
doc.add_paragraph('Application: Sensor Simulation | Framework: FastAPI with SQLAlchemy')

# Overview
doc.add_heading('Overview', 1)
overview_table = [
    ['Application', 'Sensor Simulation'],
    ['Version', '1.0.0'],
    ['Framework', 'FastAPI with SQLAlchemy async ORM'],
    ['Database', 'PostgreSQL with TimescaleDB extension'],
    ['WebServer', 'Uvicorn'],
]
add_table_to_doc(doc, 'System Information', ['Component', 'Details'], overview_table)

# REST ENDPOINTS
doc.add_heading('1. REST ENDPOINTS', 1)

doc.add_heading('Sensors Resource', 2)
sensors_endpoints = [
    ['GET', '/api/v1/sensors/summary', 'get_sensor_summary()', 'None', 'SensorSummaryOut', '200', 'Returns total, active, inactive, error sensor counts'],
    ['GET', '/api/v1/sensors', 'get_all_sensors()', 'None', 'list[SensorOut]', '200', 'Get all sensors ordered by creation date'],
    ['POST', '/api/v1/sensors', 'create_sensor()', 'SensorCreate', 'SensorOut', '201', 'Register a new sensor (default status: offline)'],
    ['GET', '/api/v1/sensors/{sensor_id}', 'get_sensor()', 'None', 'SensorOut', '200, 404', 'Get single sensor by ID'],
    ['PUT', '/api/v1/sensors/{sensor_id}', 'update_sensor()', 'SensorUpdate', 'SensorOut', '200, 404', 'Update sensor location and coverage radius'],
]
add_table_to_doc(doc, 'Sensor Endpoints', ['HTTP Method', 'Route', 'Handler', 'Request', 'Response', 'Status', 'Description'], sensors_endpoints)

doc.add_heading('Threats Resource', 2)
threats_endpoints = [
    ['GET', '/api/v1/threats/summary', 'get_threat_summary()', 'None', 'ThreatSummaryOut', '200', 'Returns total threats, high severity count, active sensor count'],
    ['GET', '/api/v1/threats', 'get_threats()', 'Query params', 'PagedThreats', '200', 'Get cursor-paginated threats with filtering'],
]
add_table_to_doc(doc, 'Threat Endpoints', ['HTTP Method', 'Route', 'Handler', 'Request', 'Response', 'Status', 'Description'], threats_endpoints)

doc.add_heading('Users Resource', 2)
users_endpoints = [
    ['POST', '/api/v1/users', 'create_user()', 'UserCreate', 'UserOut', '201', 'Create new user (default role: admin)'],
    ['GET', '/api/v1/users/{user_id}', 'get_user()', 'None', 'UserOut', '200, 404', 'Get user by ID'],
    ['PUT', '/api/v1/users/{user_id}', 'update_user()', 'UserUpdate', 'UserOut', '200, 404', 'Update username or email'],
    ['PUT', '/api/v1/users/{user_id}/change-password', 'change_password()', 'PasswordChange', 'dict', '200, 404, 400', 'Change user password with verification'],
]
add_table_to_doc(doc, 'User Endpoints', ['HTTP Method', 'Route', 'Handler', 'Request', 'Response', 'Status', 'Description'], users_endpoints)

doc.add_heading('Analytics Resource', 2)
analytics_endpoints = [
    ['GET', '/api/v1/analytics/threat-timeline', 'threat_timeline()', 'Query params', 'ThreatTimelineOut', '200', 'Threat counts grouped into time buckets'],
    ['GET', '/api/v1/analytics/threats-per-sensor', 'threats_per_sensor()', 'Query params', 'ThreatsPerSensorOut', '200', 'Threat counts grouped by sensor'],
    ['GET', '/api/v1/analytics/severity-breakdown', 'severity_breakdown()', 'Query params', 'SeverityBreakdownOut', '200', 'Threat counts grouped by severity'],
]
add_table_to_doc(doc, 'Analytics Endpoints', ['HTTP Method', 'Route', 'Handler', 'Request', 'Response', 'Status', 'Description'], analytics_endpoints)

# WEBSOCKET
doc.add_heading('2. WEBSOCKET ENDPOINTS', 1)
doc.add_paragraph('WebSocket Connection: /ws')
doc.add_paragraph('Message Format: JSON with {"event": string, "data": object}')
doc.add_paragraph('Event Types: auth_confirmed, alert_new, alert_acknowledged')

ws_table = [
    ['auth_confirmed', 'Sent on successful connection', '{"message": "Connected to Sensor Simulation", "status": "live"}'],
    ['alert_new', 'Sent when new threat detected', '{alert_data_object}'],
    ['alert_acknowledged', 'Sent when threat acknowledged', 'Threat ID in data'],
]
add_table_to_doc(doc, 'WebSocket Events', ['Event', 'Description', 'Data'], ws_table)

# DATABASE MODELS
doc.add_heading('3. DATABASE MODELS', 1)

doc.add_heading('Sensor Model', 2)
sensor_fields = [
    ['sensor_id', 'str', 'Primary Key', 'Unique identifier'],
    ['sensor_type', 'SensorType enum', 'NOT NULL', 'radar or lidar'],
    ['status', 'SensorStatus enum', 'NOT NULL, Default: inactive', 'active, inactive, error'],
    ['lat', 'float', 'NOT NULL, [-90,90]', 'Latitude coordinate'],
    ['lng', 'float', 'NOT NULL, [-180,180]', 'Longitude coordinate'],
    ['location', 'str', 'NOT NULL', 'Location name'],
    ['coverage_radius_m', 'float', 'NOT NULL, Default: 50.0', 'Coverage radius in meters'],
    ['last_ping', 'datetime', 'Nullable', 'Last heartbeat timestamp'],
    ['created_at', 'datetime', 'NOT NULL, Server Default', 'Creation timestamp'],
]
add_table_to_doc(doc, 'Sensor Fields', ['Field', 'Type', 'Constraints', 'Description'], sensor_fields)

doc.add_heading('ThreatLog Model', 2)
threat_fields = [
    ['alert_id', 'str', 'Primary Key, Default: UUID', 'Unique threat ID'],
    ['sensor_id', 'str', 'Foreign Key, Indexed', 'Reference to detecting sensor'],
    ['sensor_type', 'str', 'NOT NULL', 'Type of sensor'],
    ['threat_type', 'str', 'NOT NULL', 'Classification of threat'],
    ['confidence', 'float', 'NOT NULL', 'Confidence score (0.0-1.0)'],
    ['severity', 'ThreatSeverity enum', 'NOT NULL', 'low, med, high'],
    ['timestamp', 'datetime', 'NOT NULL, Indexed, TimescaleDB hypertable', 'Detection time'],
]
add_table_to_doc(doc, 'ThreatLog Fields', ['Field', 'Type', 'Constraints', 'Description'], threat_fields)

doc.add_heading('User Model', 2)
user_fields = [
    ['user_id', 'str', 'Primary Key, Default: UUID', 'Unique identifier'],
    ['username', 'str', 'NOT NULL, Unique', 'Login name'],
    ['email', 'str', 'NOT NULL, Unique', 'Email address'],
    ['password_hash', 'str', 'NOT NULL', 'Hashed password'],
    ['role', 'str', 'NOT NULL, Default: admin', 'User role'],
]
add_table_to_doc(doc, 'User Fields', ['Field', 'Type', 'Constraints', 'Description'], user_fields)

doc.add_heading('RadarReading Model', 2)
doc.add_paragraph('Purpose: Time-series radar sensor measurements (TimescaleDB hypertable)')
radar_fields = [
    ['id', 'str', 'Primary Key', 'Reading ID'],
    ['sensor_id', 'str', 'Foreign Key, Indexed', 'Radar sensor reference'],
    ['timestamp', 'datetime', 'Hypertable dimension', 'Reading timestamp'],
    ['status', 'ReadingStatus enum', 'Default: OK', 'OK or Threat'],
    ['range_m', 'float', 'NOT NULL', 'Distance to target (m)'],
    ['azimuth_deg', 'float', 'NOT NULL', 'Horizontal angle'],
    ['elevation_deg', 'float', 'NOT NULL', 'Vertical angle'],
    ['radial_velocity_mps', 'float', 'NOT NULL', 'Velocity toward/away'],
    ['rcs_dbsm', 'float', 'NOT NULL', 'Radar cross-section'],
    ['snr_db', 'float', 'NOT NULL', 'Signal-to-noise ratio'],
]
add_table_to_doc(doc, 'RadarReading Fields', ['Field', 'Type', 'Constraints', 'Description'], radar_fields)

doc.add_heading('LidarReading Model', 2)
doc.add_paragraph('Purpose: Time-series LIDAR sensor measurements (TimescaleDB hypertable)')
lidar_fields = [
    ['id', 'str', 'Primary Key', 'Reading ID'],
    ['sensor_id', 'str', 'Foreign Key, Indexed', 'LIDAR sensor reference'],
    ['timestamp', 'datetime', 'Hypertable dimension', 'Reading timestamp'],
    ['status', 'ReadingStatus enum', 'Default: OK', 'OK or Threat'],
    ['bbox_x_min/max', 'float', 'Bounding box', 'X coordinates (m)'],
    ['bbox_y_min/max', 'float', 'Bounding box', 'Y coordinates (m)'],
    ['bbox_z_min/max', 'float', 'Bounding box', 'Z coordinates (m)'],
    ['centroid_x/y/z', 'float', 'Centroid position', 'Center coordinates (m)'],
    ['point_count', 'int', 'NOT NULL', 'Points in cloud'],
    ['intensity_avg', 'float', 'NOT NULL', 'Average intensity'],
    ['velocity_mps', 'float', 'NOT NULL', 'Velocity (m/s)'],
    ['aspect_ratio', 'float', 'NOT NULL', 'Bbox aspect ratio'],
    ['point_density_ppm2', 'float', 'NOT NULL', 'Points per m²'],
]
add_table_to_doc(doc, 'LidarReading Fields', ['Field', 'Type', 'Constraints', 'Description'], lidar_fields)

# SERVICES
doc.add_heading('4. SERVICES', 1)

doc.add_heading('SensorService', 2)
sensor_svc = [
    ['get_all_sensors(db)', 'list[SensorOut]', 'GET /api/v1/sensors'],
    ['get_sensor(sensor_id, db)', 'SensorOut', 'GET /api/v1/sensors/{id}'],
    ['create_sensor(data, db)', 'SensorOut', 'POST /api/v1/sensors'],
    ['update_sensor(sensor_id, data, db)', 'SensorOut', 'PUT /api/v1/sensors/{id}'],
    ['get_sensor_summary(db)', 'SensorSummaryOut', 'GET /api/v1/sensors/summary'],
]
add_table_to_doc(doc, 'SensorService Methods', ['Function', 'Returns', 'Used By'], sensor_svc)

doc.add_heading('ThreatService', 2)
threat_svc = [
    ['get_threats(filters, db)', 'PagedThreats', 'GET /api/v1/threats (cursor pagination)'],
    ['get_threat_summary(db)', 'ThreatSummaryOut', 'GET /api/v1/threats/summary'],
    ['push_alert(alert_data)', 'None', 'WebSocket broadcast'],
]
add_table_to_doc(doc, 'ThreatService Methods', ['Function', 'Returns', 'Used By'], threat_svc)

doc.add_heading('AnalyticsService', 2)
analytics_svc = [
    ['get_threat_timeline(filters, bucket_by, db)', 'ThreatTimelineOut', 'GET /api/v1/analytics/threat-timeline'],
    ['get_threats_per_sensor(filters, db)', 'ThreatsPerSensorOut', 'GET /api/v1/analytics/threats-per-sensor'],
    ['get_severity_breakdown(filters, db)', 'SeverityBreakdownOut', 'GET /api/v1/analytics/severity-breakdown'],
]
add_table_to_doc(doc, 'AnalyticsService Methods', ['Function', 'Returns', 'Used By'], analytics_svc)

doc.add_heading('UserService', 2)
user_svc = [
    ['create_user(data, db)', 'UserOut', 'POST /api/v1/users'],
    ['get_user(user_id, db)', 'UserOut', 'GET /api/v1/users/{id}'],
    ['update_user(user_id, data, db)', 'UserOut', 'PUT /api/v1/users/{id}'],
    ['change_password(user_id, old_pwd, new_pwd, db)', 'dict', 'PUT /api/v1/users/{id}/change-password'],
]
add_table_to_doc(doc, 'UserService Methods', ['Function', 'Returns', 'Used By'], user_svc)

# AUTHENTICATION & CONFIGURATION
doc.add_heading('5. AUTHENTICATION & CONFIGURATION', 1)

doc.add_heading('Authentication Status', 2)
doc.add_paragraph('⚠️ NOT IMPLEMENTED — All endpoints are currently public')
doc.add_paragraph('Planned: JWT tokens, bcrypt hashing, RBAC')

doc.add_heading('CORS Configuration', 2)
doc.add_paragraph('Status: ENABLED with permissive policy (development)')
doc.add_paragraph('Allow Origins: * (all)')
doc.add_paragraph('Allow Methods: * (all)')
doc.add_paragraph('Allow Headers: * (all)')
doc.add_paragraph('Allow Credentials: True')

doc.add_heading('Database Configuration', 2)
db_config = [
    ['Engine', 'PostgreSQL + asyncpg'],
    ['Host', 'localhost (default)'],
    ['Port', '5432'],
    ['Database', 'sensor_simulation'],
    ['User', 'sensor'],
    ['Pool Size', '10 minimum, +20 overflow'],
    ['Pool Pre-Ping', 'Enabled'],
]
add_table_to_doc(doc, 'Database Settings', ['Setting', 'Value'], db_config)

doc.add_heading('TimescaleDB Setup', 2)
doc.add_paragraph('Hypertables (auto-created):')
doc.add_paragraph('• radar_readings — partitioned by timestamp')
doc.add_paragraph('• lidar_readings — partitioned by timestamp')

doc.add_heading('Server Configuration', 2)
doc.add_paragraph('Framework: FastAPI')
doc.add_paragraph('ASGI Server: Uvicorn')
doc.add_paragraph('Documentation URLs:')
doc.add_paragraph('  • Swagger UI: /api/docs')
doc.add_paragraph('  • ReDoc: /api/redoc')
doc.add_paragraph('  • OpenAPI Schema: /api/openapi.json')

# ERROR RESPONSES
doc.add_heading('6. ERROR RESPONSES', 1)

error_codes = [
    ['200', 'OK', 'Successful GET/PUT/DELETE'],
    ['201', 'Created', 'Resource created (POST)'],
    ['400', 'Bad Request', 'Invalid input parameters'],
    ['404', 'Not Found', 'Resource not found'],
    ['409', 'Conflict', 'Resource conflict (duplicate ID, already acknowledged)'],
    ['500', 'Internal Server Error', 'Server/database error'],
]
add_table_to_doc(doc, 'HTTP Status Codes', ['Code', 'Status', 'When Used'], error_codes)

doc.add_paragraph('Standard error response format:')
doc.add_paragraph('{ "detail": "Error message describing the issue" }')

# DATA VALIDATION RULES
doc.add_heading('7. DATA VALIDATION RULES', 1)

doc.add_heading('Sensor Validation', 2)
doc.add_paragraph('• sensor_id: Required, unique string')
doc.add_paragraph('• sensor_type: Required, enum ["radar", "lidar"]')
doc.add_paragraph('• lat: Required, range [-90, 90]')
doc.add_paragraph('• lng: Required, range [-180, 180]')
doc.add_paragraph('• location: Required, non-empty string')
doc.add_paragraph('• coverage_radius_m: Default 50.0, must be > 0')

doc.add_heading('Threat Validation', 2)
doc.add_paragraph('• sensor_id: Required, must reference existing sensor')
doc.add_paragraph('• threat_type: Required, non-empty string')
doc.add_paragraph('• confidence: Required, float (0.0-1.0)')
doc.add_paragraph('• severity: Required, enum ["low", "med", "high"]')
doc.add_paragraph('• timestamp: Auto-set to server NOW')

doc.add_heading('User Validation', 2)
doc.add_paragraph('• username: Required, unique string')
doc.add_paragraph('• email: Required, unique valid email')
doc.add_paragraph('• password: Required (stored as plaintext, bcrypt planned)')
doc.add_paragraph('• role: Auto-set to "admin" on creation')

doc.add_heading('Pagination Validation', 2)
doc.add_paragraph('• page_size: Default 20, recommended max 100')
doc.add_paragraph('• cursor: Optional base64-encoded {timestamp, alert_id} JSON')

doc.add_heading('Date Range Validation', 2)
doc.add_paragraph('• from_dt: Optional ISO 8601 UTC, defaults to 7 days ago')
doc.add_paragraph('• to_dt: Optional ISO 8601 UTC, defaults to NOW')
doc.add_paragraph('• bucket_by: Enum ["minute", "hour", "day"], default "hour"')

# SPECIAL FEATURES
doc.add_heading('8. SPECIAL FEATURES & EXTENSIONS', 1)

doc.add_heading('Time-Series Optimization (TimescaleDB)', 2)
doc.add_paragraph('Benefit: Optimized compression for high-volume sensor readings')
doc.add_paragraph('Tables: radar_readings, lidar_readings with automatic compression')

doc.add_heading('Live Threat Broadcasting', 2)
doc.add_paragraph('Architecture: Publish-Subscribe via WebSocket')
doc.add_paragraph('Trigger: Detection engine pushes alerts → broadcast to all clients')
doc.add_paragraph('Result: Real-time frontend updates without polling')

doc.add_heading('Cursor-Based Pagination', 2)
doc.add_paragraph('Type: Keyset pagination using (timestamp, alert_id)')
doc.add_paragraph('Advantage: Stable pagination, handles real-time inserts')

doc.add_heading('Analytics Aggregation', 2)
doc.add_paragraph('• Span: 7-day default, customizable date range')
doc.add_paragraph('• Granularity: Minute/hour/day bucketing')
doc.add_paragraph('• Completeness: All sensors included even with zero threats')

# DEPLOYMENT
doc.add_heading('9. DEPLOYMENT & REQUIREMENTS', 1)

doc.add_heading('Environment Variables (.env)', 2)
doc.add_paragraph('DB_USER=sensor')
doc.add_paragraph('DB_PASSWORD=sensor_pass')
doc.add_paragraph('DB_HOST=localhost')
doc.add_paragraph('DB_PORT=5432')
doc.add_paragraph('DB_NAME=sensor_simulation')
doc.add_paragraph('FASTAPI_DEBUG=True')

doc.add_heading('Dependencies', 2)
doc.add_paragraph('fastapi')
doc.add_paragraph('uvicorn[standard]')
doc.add_paragraph('sqlalchemy[asyncio]')
doc.add_paragraph('asyncpg')
doc.add_paragraph('pydantic')

doc.add_heading('Database Setup', 2)
doc.add_paragraph('1. Install PostgreSQL 14+ and TimescaleDB extension')
doc.add_paragraph('2. Create database: createdb sensor_simulation')
doc.add_paragraph('3. Run application — tables auto-created on startup')

doc.add_heading('Running the API', 2)
doc.add_paragraph('cd Backend')
doc.add_paragraph('uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload')
doc.add_paragraph('')
doc.add_paragraph('Access documentation: http://localhost:8000/api/docs')

# Footer
doc.add_page_break()
doc.add_paragraph('').paragraph_format.space_after = Pt(12)
footer = doc.add_paragraph('Document: Backend API Documentation')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer = doc.add_paragraph('Version: 1.0.0 | Generated: March 30, 2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer = doc.add_paragraph('Sensor Simulation Project')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\ShreyaPradeepAnkalgi\Sensor-Simulation-Python-Script\Backend_API_Documentation.docx'
doc.save(output_path)
print(f"✓ Document created: {output_path}")
